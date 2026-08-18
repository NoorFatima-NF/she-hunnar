import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x78, 0x35, 0x0F) # Amber 900
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46) # Emerald 800
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
    return h

def create_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1E293B") # Dark slate
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table

def generate_doc():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # -------------------------------------------------------------
    # TITLE & HEADER
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("SHE HUNNAR MARKETPLACE")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x78, 0x35, 0x0F) # Amber 900

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Complete Database Architecture, Entity Relationship & Data Dictionary Report")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Project", "She Hunnar Multi-Vendor"), ("Engine", "PostgreSQL / Supabase / Prisma"), ("Safepay Status", "Sandbox 3D-Secure"), ("Version", "1.0 Production")],
        [("Architecture", "Multi-Vendor Split Escrow"), ("Total Tables", "21 Relational Tables"), ("Author", "Lead Database Architect"), ("Date", "August 2026")]
    ]
    for r_i, row in enumerate(meta_data):
        for c_i, (k, v) in enumerate(row):
            cell = meta_table.rows[r_i].cells[c_i]
            cell.text = f"{k}: {v}"
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            p.runs[0].font.bold = (k in ["Project", "Total Tables", "Safepay Status"])

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    add_styled_heading(doc, "1. Executive Summary & Architecture Overview", level=1)
    p = doc.add_paragraph(
        "She Hunnar is a specialized handmade jewelry and artisan craft marketplace designed to empower independent craft makers and female artisans across Pakistan. The database architecture is engineered around high performance, strict referential integrity, automated multi-vendor order fulfillment, and secure escrow payment handling."
    )
    p.paragraph_format.line_spacing = 1.15

    doc.add_paragraph(
        "Core Architectural Highlights:\n"
        "• Multi-Vendor Sub-Order Splitting: When a customer orders items from multiple artisan shops, the system records one Master Order and automatically splits it into individual Seller Sub-Orders for separate tracking, courier generation, and shipping calculations.\n"
        "• Safepay Sandbox Integration: Fully integrated with Safepay's official Hosted Checkout API, storing session tracker tokens, HMAC-SHA256 verification signatures, and payment verification timestamps.\n"
        "• Bespoke Customization Engine: Supports custom engraved names, initials, font choices, birthstones, and buyer notes.\n"
        "• Artisan Studio Governance: Comprehensive verification with CNIC snapshots, pickup addresses, commission deductions, and automated payout requests."
    )

    # -------------------------------------------------------------
    # SECTION 2: ENTITY RELATIONSHIP (ER) DIAGRAM & MAPPINGS
    # -------------------------------------------------------------
    add_styled_heading(doc, "2. Entity Relationship (ER) Diagram & Table Relationships", level=1)
    doc.add_paragraph(
        "Below is the complete architectural entity relationship mapping illustrating foreign key hierarchies, cardinalities, and cascades across all subsystems:"
    )

    er_headers = ["Parent Entity", "Relationship", "Child Entity", "Cardinality", "Foreign Key / Cascade Rule"]
    er_data = [
        ["users", "owns studio", "seller_profiles", "1 : 0..1", "seller_profiles.user_id -> users.id (CASCADE)"],
        ["users", "saves address", "addresses", "1 : N", "addresses.user_id -> users.id (CASCADE)"],
        ["users", "places", "master_orders", "1 : N", "master_orders.customer_id -> users.id (SET NULL)"],
        ["users", "adds item", "cart_items", "1 : N", "cart_items.user_id -> users.id (CASCADE)"],
        ["users", "saves favorite", "wishlist_items", "1 : N", "wishlist_items.user_id -> users.id (CASCADE)"],
        ["users", "submits review", "product_reviews", "1 : N", "product_reviews.customer_id -> users.id (SET NULL)"],
        ["seller_profiles", "publishes", "products", "1 : N", "products.seller_id -> seller_profiles.id (CASCADE)"],
        ["categories", "classifies", "products", "1 : N", "products.category_id -> categories.id (SET NULL)"],
        ["products", "has sizes/colors", "product_variants", "1 : N", "product_variants.product_id -> products.id (CASCADE)"],
        ["products", "has custom options", "product_customization_configs", "1 : 1", "product_customization_configs.product_id -> products.id (CASCADE)"],
        ["master_orders", "splits into", "seller_orders", "1 : N", "seller_orders.master_order_id -> master_orders.id (CASCADE)"],
        ["seller_profiles", "fulfills", "seller_orders", "1 : N", "seller_orders.seller_id -> seller_profiles.id (RESTRICT)"],
        ["seller_orders", "contains items", "order_items", "1 : N", "order_items.seller_order_id -> seller_orders.id (CASCADE)"],
        ["products", "ordered in", "order_items", "1 : N", "order_items.product_id -> products.id (SET NULL)"],
        ["users & sellers", "creates chat", "conversations", "M : N", "conversations.customer_id & seller_id (UNIQUE)"],
        ["conversations", "contains", "messages", "1 : N", "messages.conversation_id -> conversations.id (CASCADE)"],
        ["seller_profiles", "creates promo", "coupons", "1 : N", "coupons.applicable_seller_id -> seller_profiles.id (CASCADE)"],
        ["seller_profiles", "withdraws funds", "seller_payouts", "1 : N", "seller_payouts.seller_id -> seller_profiles.id (CASCADE)"],
        ["master_orders", "has dispute", "return_requests", "1 : N", "return_requests.master_order_id -> master_orders.id (CASCADE)"],
        ["users", "requests custom piece", "custom_jewelry_requests", "1 : N", "custom_jewelry_requests.customer_id -> users.id (SET NULL)"]
    ]
    create_table(doc, er_headers, er_data, [1.3, 1.2, 1.5, 0.9, 2.3])

    # -------------------------------------------------------------
    # SECTION 3: COMPREHENSIVE DATA DICTIONARY (21 TABLES)
    # -------------------------------------------------------------
    add_styled_heading(doc, "3. Comprehensive Table-by-Table Data Dictionary", level=1)
    doc.add_paragraph(
        "This section details every table, its precise business function, all column data types, constraints, and operational logic."
    )

    # List of tables and detailed columns
    tables_dict = [
        {
            "name": "1. users",
            "desc": "Stores all authenticated accounts including retail customers, artisan sellers, and super admins. Supports OAuth Google sign-in and password hashing.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY, DEFAULT uuid_generate_v4()", "Unique user identifier"],
                ["name", "VARCHAR(255)", "NOT NULL", "Full display name"],
                ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Login email address"],
                ["password_hash", "VARCHAR(255)", "NULLABLE", "Bcrypt/Argon2 password hash"],
                ["phone", "VARCHAR(50)", "NULLABLE", "Contact phone number"],
                ["role", "ENUM (user_role)", "DEFAULT 'customer'", "Role: customer, seller, admin"],
                ["avatar_url", "TEXT", "NULLABLE", "Profile picture URL"],
                ["google_id", "VARCHAR(255)", "NULLABLE", "Google OAuth identifier"],
                ["is_active", "BOOLEAN", "DEFAULT TRUE", "Account active/banned status"],
                ["created_at / updated_at", "TIMESTAMP WITH TIME ZONE", "DEFAULT CURRENT_TIMESTAMP", "Record audit timestamps"]
            ]
        },
        {
            "name": "2. addresses",
            "desc": "Maintains delivery shipping destinations for customers with default address selection.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Address record identifier"],
                ["user_id", "UUID", "FK -> users(id) ON DELETE CASCADE", "Associated user account"],
                ["full_name", "VARCHAR(255)", "NOT NULL", "Recipient full name"],
                ["phone", "VARCHAR(50)", "NOT NULL", "Courier SMS phone number"],
                ["address_line", "TEXT", "NOT NULL", "House/Street/Building complete address"],
                ["city", "VARCHAR(100)", "NOT NULL", "City (Lahore, Karachi, Islamabad, etc.)"],
                ["province", "VARCHAR(100)", "NOT NULL", "Punjab, Sindh, KP, Balochistan, etc."],
                ["postal_code", "VARCHAR(20)", "NOT NULL", "Postal ZIP code"],
                ["is_default", "BOOLEAN", "DEFAULT FALSE", "Flag for primary delivery address"]
            ]
        },
        {
            "name": "3. seller_profiles",
            "desc": "Contains artisan maker studio information, verification documents, commission rate, shipping fee rules, and bank withdrawal account details.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Seller profile identifier"],
                ["user_id", "UUID", "UNIQUE, FK -> users(id)", "User account owning the studio"],
                ["shop_name", "VARCHAR(255)", "NOT NULL", "Public artisan studio shop name"],
                ["slug", "VARCHAR(255)", "UNIQUE, NOT NULL", "URL shop slug (e.g. noor-jewelry-studio)"],
                ["specialization", "VARCHAR(255)", "NOT NULL", "Craft type (Silver, Resin, Calligraphy, etc.)"],
                ["rating / review_count", "NUMERIC(3,2) / INT", "DEFAULT 5.00 / 0", "Aggregated rating & total reviews"],
                ["verification_status", "ENUM", "DEFAULT 'pending'", "pending, approved, rejected, suspended"],
                ["cnic_number / cnic_name", "VARCHAR", "NULLABLE", "Government ID verification info"],
                ["cnic_front_url / back_url", "TEXT", "NULLABLE", "Uploaded identity document images"],
                ["commission_rate", "NUMERIC(5,2)", "DEFAULT 10.00", "Platform commission percentage (e.g. 10%)"],
                ["shipping_fee", "NUMERIC(10,2)", "DEFAULT 200.00", "Default shipping fee in PKR"],
                ["free_shipping_threshold", "NUMERIC(10,2)", "NULLABLE", "Free shipping threshold in PKR"],
                ["payout_method / account_details", "VARCHAR / TEXT", "DEFAULT 'Bank Transfer'", "IBAN, Easypaisa, or JazzCash details"],
                ["instagram / facebook / whatsapp", "VARCHAR", "NULLABLE", "Artisan social contact links"]
            ]
        },
        {
            "name": "4. categories",
            "desc": "Stores product classification taxonomy across the marketplace.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Category ID"],
                ["name", "VARCHAR(100)", "UNIQUE, NOT NULL", "Category Name (Jewelry, Bags, Candles, etc.)"],
                ["slug", "VARCHAR(100)", "UNIQUE, NOT NULL", "URL friendly slug"],
                ["description", "TEXT", "NULLABLE", "Category SEO overview"],
                ["image_url", "TEXT", "NULLABLE", "Cover thumbnail image"],
                ["display_order", "INT", "DEFAULT 0", "Homepage priority order"],
                ["is_active", "BOOLEAN", "DEFAULT TRUE", "Category visibility"]
            ]
        },
        {
            "name": "5. products",
            "desc": "Main handmade jewelry & craft catalog with materials, stock, pricing, production timeline, and care guides.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Product ID"],
                ["seller_id", "UUID", "FK -> seller_profiles(id)", "Maker studio owner"],
                ["category_name", "VARCHAR(100)", "NOT NULL", "Jewelry, Bags, Calligraphy, Home Decor, etc."],
                ["title / slug", "VARCHAR", "UNIQUE slug", "Product title and URL slug"],
                ["material", "VARCHAR(255)", "NOT NULL", "925 Sterling Silver, Freshwater Pearl, etc."],
                ["metal_type / stone_type", "VARCHAR", "NULLABLE", "Gold plated, Zircon, Raw Emerald, etc."],
                ["price / original_price", "NUMERIC(10,2)", "NOT NULL", "Selling price and original price in PKR"],
                ["stock / sku", "INT / VARCHAR", "NOT NULL, UNIQUE sku", "Available inventory & SKU identifier"],
                ["images", "TEXT[]", "DEFAULT '{}'", "Array of high-resolution product image URLs"],
                ["production_time_days", "INT", "DEFAULT 3", "Days required for handmade fabrication"],
                ["status", "ENUM", "DEFAULT 'published'", "draft, pending, published, rejected, out_of_stock"],
                ["care_instructions / process", "TEXT", "NULLABLE", "Artisan process & maintenance guide"]
            ]
        },
        {
            "name": "6. product_variants",
            "desc": "Handles size selections (Ring size 6/7/8, Chain length 18/20 inch, color variations) with dedicated price & inventory tracking.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Variant ID"],
                ["product_id", "UUID", "FK -> products(id) ON DELETE CASCADE", "Parent product"],
                ["name", "VARCHAR(255)", "NOT NULL", "Variant title (e.g. Silver / Size 7)"],
                ["price", "NUMERIC(10,2)", "NOT NULL", "Variant specific price in PKR"],
                ["stock", "INT", "DEFAULT 0", "Inventory for this specific variant"],
                ["sku", "VARCHAR(100)", "NOT NULL", "Variant SKU code"]
            ]
        },
        {
            "name": "7. product_customization_configs",
            "desc": "Defines what custom options a buyer can select (engraved names, fonts, birthstones, custom notes).",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Config ID"],
                ["product_id", "UUID", "UNIQUE, FK -> products(id)", "Associated customizable product"],
                ["allow_text / text_label", "BOOLEAN / VARCHAR", "DEFAULT FALSE", "Name / Initials engraving field"],
                ["max_characters", "INT", "DEFAULT 20", "Max allowed characters for engraving"],
                ["allow_font_selection / fonts", "BOOLEAN / TEXT[]", "DEFAULT FALSE", "Supported font styles list"],
                ["allow_stone_selection / stones", "BOOLEAN / TEXT[]", "DEFAULT FALSE", "Birthstone / crystal options list"],
                ["allow_reference_upload", "BOOLEAN", "DEFAULT FALSE", "Allows buyer to upload reference photo"],
                ["allow_note / note_label", "BOOLEAN / VARCHAR", "DEFAULT FALSE", "Gift note / custom instructions"]
            ]
        },
        {
            "name": "8. cart_items",
            "desc": "Persistent multi-vendor shopping cart for registered users and guest visitors.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Cart item ID"],
                ["user_id", "UUID", "NULLABLE, FK -> users(id)", "Logged-in customer account"],
                ["session_token", "VARCHAR(255)", "NULLABLE", "Guest visitor session cookie token"],
                ["product_id", "UUID", "FK -> products(id)", "Selected product"],
                ["variant_id", "UUID", "NULLABLE, FK -> product_variants", "Selected size/variation"],
                ["quantity", "INT", "NOT NULL, DEFAULT 1", "Quantity selected"],
                ["customization_data", "JSONB", "NULLABLE", "Buyer customized text, font & note snapshot"]
            ]
        },
        {
            "name": "9. wishlist_items",
            "desc": "Saved favorite items for customer accounts.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Wishlist ID"],
                ["user_id", "UUID", "FK -> users(id)", "Customer ID"],
                ["product_id", "UUID", "FK -> products(id)", "Favorited product ID"],
                ["created_at", "TIMESTAMP WITH TIME ZONE", "DEFAULT NOW()", "Date saved to wishlist"]
            ]
        },
        {
            "name": "10. master_orders",
            "desc": "Primary customer invoice storing grand totals, shipping address snapshot, Safepay gateway tokens, signatures, and payment confirmation status.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Master Order ID"],
                ["order_number", "VARCHAR(50)", "UNIQUE, NOT NULL", "Customer visible ID (e.g. ORD-89201)"],
                ["customer_id", "UUID", "NULLABLE, FK -> users(id)", "Customer account ID"],
                ["customer_name / email / phone", "VARCHAR", "NOT NULL", "Delivery recipient contact info"],
                ["shipping_address_line / city / postal", "TEXT / VARCHAR", "NOT NULL", "Delivery address snapshot"],
                ["payment_method", "ENUM", "NOT NULL", "COD, Bank Transfer, Easypaisa, JazzCash, Online Card"],
                ["payment_status", "ENUM", "DEFAULT 'Pending'", "Pending, Paid, Failed, Refunded"],
                ["safepay_tracker", "VARCHAR(255)", "NULLABLE", "Safepay Tracker Session Token (track_xxx)"],
                ["safepay_signature", "TEXT", "NULLABLE", "Safepay HMAC-SHA256 signature hash"],
                ["transaction_ref", "VARCHAR(255)", "NULLABLE", "Payment gateway transaction reference"],
                ["payment_verified_at", "TIMESTAMP WITH TIME ZONE", "NULLABLE", "Timestamp when payment was verified"],
                ["subtotal / total_shipping / grand_total", "NUMERIC(10,2)", "NOT NULL", "Financial breakdown in PKR"],
                ["coupon_code / discount_amount", "VARCHAR / NUMERIC", "NULLABLE", "Applied discount coupon details"],
                ["master_status", "ENUM", "DEFAULT 'Pending'", "Overall multi-vendor fulfillment status"]
            ]
        },
        {
            "name": "11. seller_orders (Sub-Orders)",
            "desc": "Multi-vendor split sub-orders generated per artisan maker studio. Controls independent courier dispatch, tracking numbers, and seller payouts.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Seller Sub-Order ID"],
                ["sub_order_number", "VARCHAR(50)", "UNIQUE, NOT NULL", "Visible Sub-Order code (e.g. SO-89201-A)"],
                ["master_order_id", "UUID", "FK -> master_orders(id)", "Parent master invoice"],
                ["seller_id", "UUID", "FK -> seller_profiles(id)", "Maker studio fulfilling this portion"],
                ["subtotal / shipping_fee / total_amount", "NUMERIC(10,2)", "NOT NULL", "Subtotal & studio shipping in PKR"],
                ["status", "ENUM", "DEFAULT 'Pending'", "Pending, Confirmed, Processing, Shipped, Delivered, etc."],
                ["courier_name", "VARCHAR(100)", "NULLABLE", "Courier service (TCS, Leopards, PostEx, Trax)"],
                ["tracking_number / tracking_url", "VARCHAR / TEXT", "NULLABLE", "Courier tracking code & tracking link"],
                ["shipped_at / delivered_at", "TIMESTAMP WITH TIME ZONE", "NULLABLE", "Dispatch & delivery timestamps"]
            ]
        },
        {
            "name": "12. order_items",
            "desc": "Line item snapshot for each purchased product within a seller's sub-order.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Line item ID"],
                ["seller_order_id", "UUID", "FK -> seller_orders(id)", "Associated seller sub-order"],
                ["product_id", "UUID", "NULLABLE, FK -> products(id)", "Purchased product"],
                ["product_title / product_image", "VARCHAR / TEXT", "NOT NULL", "Snapshot of title & image at purchase"],
                ["variant_name", "VARCHAR(255)", "NULLABLE", "Selected size / chain variation name"],
                ["price / quantity", "NUMERIC(10,2) / INT", "NOT NULL", "Price per unit & quantity ordered"],
                ["customization_data", "JSONB", "NULLABLE", "Snapshot of custom engraving / note data"]
            ]
        },
        {
            "name": "13. product_reviews",
            "desc": "Customer feedback, 1-5 star ratings, buyer photo uploads, and verified purchase flags with artisan replies.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Review ID"],
                ["product_id", "UUID", "FK -> products(id)", "Reviewed product"],
                ["customer_id", "UUID", "NULLABLE, FK -> users(id)", "Reviewer user ID"],
                ["customer_name", "VARCHAR(255)", "NOT NULL", "Reviewer display name"],
                ["rating", "INT", "NOT NULL (1 to 5)", "Rating score from 1 to 5"],
                ["review_text", "TEXT", "NOT NULL", "Customer feedback review text"],
                ["verified_purchase", "BOOLEAN", "DEFAULT TRUE", "Verified buyer badge"],
                ["images", "TEXT[]", "DEFAULT '{}'", "Customer uploaded review photos"],
                ["seller_reply / seller_reply_at", "TEXT / TIMESTAMP", "NULLABLE", "Artisan response message & date"]
            ]
        },
        {
            "name": "14. coupons",
            "desc": "Promotional discount codes supporting percentage or fixed discounts, minimum spends, and seller-specific restrictions.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Coupon ID"],
                ["code", "VARCHAR(50)", "UNIQUE, NOT NULL", "Promo code (e.g. WELCOME10, CRAFT500)"],
                ["discount_type", "ENUM", "DEFAULT 'percentage'", "percentage or fixed"],
                ["discount_value", "NUMERIC(10,2)", "NOT NULL", "Discount amount (e.g. 15% or PKR 500)"],
                ["min_spend / max_discount", "NUMERIC(10,2)", "NOT NULL / NULLABLE", "Minimum order spend & maximum cap"],
                ["expiry_date", "TIMESTAMP WITH TIME ZONE", "NOT NULL", "Expiration timestamp"],
                ["is_active", "BOOLEAN", "DEFAULT TRUE", "Active/disabled status"],
                ["applicable_seller_id", "UUID", "NULLABLE, FK -> seller_profiles", "Restricted to specific seller or storewide"]
            ]
        },
        {
            "name": "15. conversations",
            "desc": "Direct messaging rooms between customers and artisan shops.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Conversation ID"],
                ["customer_id", "UUID", "FK -> users(id)", "Customer participant"],
                ["seller_id", "UUID", "FK -> seller_profiles(id)", "Artisan studio participant"],
                ["created_at / updated_at", "TIMESTAMP WITH TIME ZONE", "DEFAULT NOW()", "Chat creation & last message date"]
            ]
        },
        {
            "name": "16. messages",
            "desc": "Individual chat messages with image attachment support and read receipts.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Message ID"],
                ["conversation_id", "UUID", "FK -> conversations(id)", "Parent chat room"],
                ["sender_id", "UUID", "FK -> users(id)", "Sender user ID"],
                ["sender_role", "ENUM", "NOT NULL", "customer or seller"],
                ["text", "TEXT", "NOT NULL", "Message body text"],
                ["attachment_url", "TEXT", "NULLABLE", "Attached image URL"],
                ["is_read", "BOOLEAN", "DEFAULT FALSE", "Read receipt status"]
            ]
        },
        {
            "name": "17. notifications",
            "desc": "Real-time alerts for customers (order shipped), sellers (new order received), and admins (payout requests).",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Notification ID"],
                ["target_role", "ENUM", "NOT NULL", "customer, seller, or admin"],
                ["target_user_id / target_seller_id", "UUID", "NULLABLE", "Target recipient ID"],
                ["title / message", "VARCHAR / TEXT", "NOT NULL", "Notification title and body"],
                ["link", "TEXT", "NULLABLE", "In-app navigation link"],
                ["is_read", "BOOLEAN", "DEFAULT FALSE", "Read status"]
            ]
        },
        {
            "name": "18. seller_payouts",
            "desc": "Records artisan studio earnings withdrawal requests, commission deductions, and payment proof.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Payout ID"],
                ["seller_id", "UUID", "FK -> seller_profiles(id)", "Artisan requesting withdrawal"],
                ["amount", "NUMERIC(10,2)", "NOT NULL", "Withdrawal amount in PKR"],
                ["period", "VARCHAR(100)", "NOT NULL", "Billing cycle period (e.g. August 2026)"],
                ["status", "ENUM", "DEFAULT 'Pending'", "Pending, Approved, Paid"],
                ["payout_method / account_details", "VARCHAR / TEXT", "NOT NULL", "Bank / Easypaisa transfer info"],
                ["requested_at / paid_at", "TIMESTAMP WITH TIME ZONE", "DEFAULT NOW() / NULL", "Request and transfer timestamps"]
            ]
        },
        {
            "name": "19. return_requests",
            "desc": "Manages 7-day handmade exchange, damage dispute, and refund requests.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Return Request ID"],
                ["master_order_id / seller_order_id", "UUID", "NOT NULL", "Associated order IDs"],
                ["seller_id / customer_id / product_id", "UUID", "NOT NULL", "Involved parties & item"],
                ["reason / description", "VARCHAR / TEXT", "NOT NULL", "Reason for return & detailed explanation"],
                ["status", "VARCHAR(50)", "DEFAULT 'Requested'", "Requested, Under Review, Approved, Rejected, Refunded"],
                ["images", "TEXT[]", "DEFAULT '{}'", "Photo proof of damage/issue"]
            ]
        },
        {
            "name": "20. custom_jewelry_requests",
            "desc": "Custom bespoke handmade jewelry quotation system where customers request custom designs with budget and material preferences.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Request ID"],
                ["customer_id", "UUID", "NULLABLE, FK -> users(id)", "Customer user ID"],
                ["customer_name / email / phone", "VARCHAR", "NOT NULL", "Contact details"],
                ["jewelry_type", "VARCHAR(100)", "NOT NULL", "Necklace, Ring, Bracelet, Calligraphy, etc."],
                ["budget_pkr", "NUMERIC(10,2)", "NOT NULL", "Customer budget in PKR"],
                ["preferred_material", "VARCHAR(255)", "NOT NULL", "925 Silver, Gold plated, Gemstones, etc."],
                ["description", "TEXT", "NOT NULL", "Design description & specifications"],
                ["reference_images", "TEXT[]", "DEFAULT '{}'", "Uploaded reference photo sketches"],
                ["preferred_seller_id", "UUID", "NULLABLE, FK -> seller_profiles", "Specific maker requested"],
                ["status", "ENUM", "DEFAULT 'Submitted'", "Submitted, In Review, Quoted, In Production, Completed"]
            ]
        },
        {
            "name": "21. platform_settings",
            "desc": "Global marketplace configurations including platform commission rates, announcement bar banners, and support details.",
            "cols": [
                ["id", "UUID", "PRIMARY KEY", "Settings ID"],
                ["commission_rate", "NUMERIC(5,2)", "DEFAULT 10.00", "Default platform commission percentage"],
                ["announcement_text", "TEXT", "NULLABLE", "Top bar promotional announcement text"],
                ["currency", "VARCHAR(10)", "DEFAULT 'PKR'", "Base platform currency"],
                ["min_payout_amount", "NUMERIC(10,2)", "DEFAULT 1000.00", "Minimum balance required for seller payout"],
                ["support_email / support_phone", "VARCHAR", "NOT NULL", "Official marketplace customer support info"],
                ["updated_at", "TIMESTAMP WITH TIME ZONE", "DEFAULT NOW()", "Last modified timestamp"]
            ]
        }
    ]

    for tbl in tables_dict:
        add_styled_heading(doc, tbl["name"], level=2)
        p_desc = doc.add_paragraph()
        p_desc.add_run("Purpose & Function: ").bold = True
        p_desc.add_run(tbl["desc"])
        p_desc.paragraph_format.space_after = Pt(6)

        tbl_headers = ["Column Name", "Data Type", "Constraints & Default", "Description / Business Logic"]
        create_table(doc, tbl_headers, tbl["cols"], [1.5, 1.4, 1.8, 2.5])

    # -------------------------------------------------------------
    # SECTION 4: STEP-BY-STEP DEPLOYMENT & SETUP GUIDE
    # -------------------------------------------------------------
    add_styled_heading(doc, "4. Database Deployment & Execution Guide", level=1)
    doc.add_paragraph(
        "You can instantly deploy this complete database schema to any modern cloud or local database using the methods below:"
    )

    doc.add_paragraph(
        "Option A: Deploy to Supabase / PostgreSQL (Recommended)\n"
        "1. Open your Supabase Dashboard or pgAdmin.\n"
        "2. Navigate to SQL Editor.\n"
        "3. Copy and run the complete script from database/schema.sql.\n"
        "4. (Optional) Run database/seed.sql to insert default categories, admin user, and sample coupons.\n\n"
        "Option B: Deploy via Prisma ORM\n"
        "1. Add DATABASE_URL to your .env file:\n"
        "   DATABASE_URL=\"postgresql://postgres:password@localhost:5432/shehunnar_db\"\n"
        "2. Run the migration command:\n"
        "   npx prisma db push\n"
        "   npx prisma generate\n"
        "3. All TypeScript types and database tables will be generated automatically."
    )

    output_path = os.path.abspath("She_Hunnar_Database_Documentation_Report.docx")
    doc.save(output_path)
    print(f"Report generated successfully at: {output_path}")

if __name__ == "__main__":
    generate_doc()
